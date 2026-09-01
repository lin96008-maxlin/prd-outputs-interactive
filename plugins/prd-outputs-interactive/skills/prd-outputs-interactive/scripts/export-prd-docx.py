#!/usr/bin/env python3
"""将主 PRD Markdown 导出为便于评审和归档的 DOCX 派生件。"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:
    raise SystemExit("缺少 python-docx，请在具备文档依赖的工作区运行。") from exc


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")


def clean_inline(text: str) -> str:
    """移除常用 Markdown 标记，同时保留链接地址和正文信息。"""
    text = re.sub(r"!\[([^]]*)]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"(`{1,3}|\*\*|__|~~)", "", text)
    text = re.sub(r"(?<!\*)\*(?!\*)|(?<!_)_(?!_)", "", text)
    return text.replace("<br>", "\n").replace("<br/>", "\n").strip()


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell.strip().replace(r"\|", "|")) for cell in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(SEPARATOR_CELL_RE.fullmatch(cell.replace(" ", "")) for cell in cells)


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 10.5) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_text_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    for index, part in enumerate(clean_inline(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        set_run_font(paragraph.add_run(part))
    return paragraph


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) >= 2 and is_separator_row(rows[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            value = values[column_index] if column_index < len(values) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
                    run.bold = row_index == 0


def export_markdown(source: Path, output: Path) -> dict[str, object]:
    markdown = source.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    document = Document()
    style_document(document)
    document.core_properties.title = source.stem
    document.core_properties.subject = "PRD"

    paragraph_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_text_paragraph(document, " ".join(paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_text_paragraph(document, "\n".join(code_buffer), style="Intense Quote")
                code_buffer.clear()
            in_code = not in_code
            index += 1
            continue

        if in_code:
            code_buffer.append(line)
            index += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            flush_paragraph()
            document.add_heading(clean_inline(heading.group(2)), level=len(heading.group(1)))
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            rows: list[list[str]] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                rows.append(split_table_row(candidate))
                index += 1
            add_markdown_table(document, rows)
            continue

        bullet = BULLET_RE.match(line)
        if bullet:
            flush_paragraph()
            add_text_paragraph(document, bullet.group(1), style="List Bullet")
            index += 1
            continue

        numbered = NUMBER_RE.match(line)
        if numbered:
            flush_paragraph()
            add_text_paragraph(document, numbered.group(1), style="List Number")
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            add_text_paragraph(document, stripped.lstrip("> "), style="Intense Quote")
            index += 1
            continue

        if not stripped:
            flush_paragraph()
        elif re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
        else:
            paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()
    if code_buffer:
        add_text_paragraph(document, "\n".join(code_buffer), style="Intense Quote")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "status": "pass",
        "source": str(source),
        "output": str(output),
        "source_characters": len(markdown),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="将 PRD Markdown 导出为 DOCX")
    parser.add_argument("source", type=Path, help="主 PRD Markdown 文件")
    parser.add_argument("output", nargs="?", type=Path, help="输出 DOCX；默认与源文件同名")
    args = parser.parse_args()

    source = args.source.resolve()
    if not source.is_file() or source.suffix.lower() != ".md":
        parser.error("source 必须是存在的 .md 文件")
    output = (args.output or source.with_suffix(".docx")).resolve()
    if output.suffix.lower() != ".docx":
        parser.error("output 必须使用 .docx 扩展名")

    result = export_markdown(source, output)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
